#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_docx.py —— 网络项目部署文档生成器（由 network-deploy 技能调用）

用法：
    python gen_docx.py [项目模型.json] [输出.docx]
    python gen_docx.py [项目模型.json] [输出.docx] --mask-secrets   # 密码脱敏
    python gen_docx.py [项目模型.json] [输出.docx] --no-toc          # 不生成目录
    python gen_docx.py [项目模型.json] [输出.docx] --credentials 凭据清单.docx  # 明文凭据另存

说明：
    - 第一个参数为 JSON 模型路径（默认 ./项目模型.json）
    - 第二个参数为输出 docx 路径（默认 ./项目设计文档.docx）
    - 依赖 python-docx；若缺失请先安装：
        python -m pip install python-docx

JSON 模型结构（详见 references/model_schema.json）：
    兼容旧结构（meta/topology/ledgers/narrative/key_decisions），
    支持工程化新结构（devices/vlans/links/routing/security/services）：
    - 若模型含 devices/vlans/links，自动生成「IP 地址分配总表」「VLAN 规划表」「互联链路表」；
    - 若含 security.acls，自动生成「ACL 策略表」；
    - 文档含敏感信息（账号密码），--mask-secrets 时密码打码为 ***，
      明文凭据通过 --credentials 另存为独立文件（建议加密保管）。

输出文档章节：
    封面信息 → 一、网络拓扑图 → 二、设备台账与配置总表 → 三、关键决策记录
    → 四、工程化规划表（IP/VLAN/互联/ACL，模型含新结构时自动生成）
    → 五、项目详细说明
"""

import argparse
import ipaddress
import json
import os
import re
import sys

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SECRET_KEYS = ("password", "secret", "community", "enable", "credential", "token")


def ensure_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.stderr.write(
            "ERROR: 缺少 python-docx。请在运行环境的虚拟环境中安装：\n"
            "  python -m pip install python-docx\n"
        )
        sys.exit(2)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    return Document, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT


def set_cjk(run, font="宋体"):
    """让 run 同时支持中文与西文显示。"""
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)


def set_cell_shading(cell, fill):
    """给单元格设置底色（w:shd）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


def repeat_header_row(table):
    """让表头跨页重复（tblHeader）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if not table.rows:
        return
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def set_table_col_widths(table, widths_cm):
    """按给定列宽（厘米）设置表格列宽（含每格固定布局）。"""
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def add_toc(doc):
    """插入 Word 目录域（打开文档后按 F9 或确认更新域即可生成目录）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "目录将在打开文档后更新（全选 Ctrl+A → F9）"
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)
    return para


def add_page_number_footer(doc):
    """页脚插入页码（第 X 页 共 Y 页）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _fld(run, instr):
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instr
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

    r1 = p.add_run("第 ")
    _fld(p.add_run(), 'PAGE')
    r2 = p.add_run(" 页 / 共 ")
    _fld(p.add_run(), 'NUMPAGES')
    r3 = p.add_run(" 页")
    for r in (r1, r2, r3):
        set_cjk(r, "宋体")
        r.font.size = Pt(9)


def mask_secret_value(value, key=""):
    """脱敏：值含密码/密钥字样时打码为 ***。"""
    if value is None:
        return value
    s = str(value)
    kl = key.lower()
    if any(k in kl for k in SECRET_KEYS) or any(k in s.lower() for k in ("@pass", "passw", "secret", "cipher")):
        return "***"
    return s


def validate(model):
    """轻量必填校验：缺关键字段给出明确错误；key_decisions 缺失仅警告。"""
    errors = []
    if not isinstance(model, dict):
        errors.append("JSON 根必须是对象")
        return errors
    meta = model.get("meta", {})
    if not isinstance(meta, dict) or not meta.get("project_name"):
        errors.append("缺少 meta.project_name（项目名称）")
    if not model.get("topology"):
        errors.append("缺少 topology（拓扑文本）")
    if not model.get("ledgers"):
        errors.append("缺少 ledgers（台账表）")
    if not model.get("narrative"):
        errors.append("缺少 narrative（文字详述）")
    if "key_decisions" not in model:
        sys.stderr.write(
            "WARN: 模型未含 key_decisions 字段，文档将缺少『关键决策记录表』；"
            "建议按 references/model_schema.json 补上阶段3/4/5确认项。\n"
        )
    return errors


def _ip_addr(s):
    try:
        return str(ipaddress.ip_interface(str(s)).ip)
    except Exception:
        return str(s)


def build_auto_tables(model, mask):
    """从工程化结构自动生成：IP 总表 / VLAN 表 / 互联表 / ACL 表。"""
    tables = []
    devices = model.get("devices", [])
    vlans = model.get("vlans", [])
    links = model.get("links", [])
    acls = (model.get("security", {}) or {}).get("acls", [])

    # IP 地址分配总表
    rows = []
    for dev in devices:
        hn = dev.get("hostname", "?")
        for key, label in (("mgmt_ip", "管理IP"), ("loopback", "Loopback")):
            if dev.get(key):
                rows.append([hn, label, _ip_addr(dev[key]), ""])
        for itf in dev.get("interfaces", []):
            if itf.get("ip"):
                rows.append([hn, "接口 " + str(itf.get("name", "")), _ip_addr(itf["ip"]),
                             str(itf.get("desc", ""))])
    if rows:
        tables.append(("IP 地址分配总表", ["设备", "用途", "地址", "备注"], rows, [2.5, 3.5, 4.0, 6.0]))

    # VLAN 规划表
    rows = []
    for v in vlans:
        vrrp = "是 (vrid %s, 主 %s)" % (v.get("vrrp_vrid") or "-", v.get("vrrp_primary") or "-") if v.get("vrrp") else "否"
        rows.append([str(v.get("id", "")), v.get("name", ""), str(v.get("net", "")),
                     str(v.get("gw", "")), vrrp])
    if rows:
        tables.append(("VLAN 规划表", ["VLAN ID", "名称", "网段", "网关", "VRRP"], rows, [2.0, 3.0, 4.0, 3.5, 4.0]))

    # 互联链路表
    rows = []
    for lk in links:
        rows.append([str(lk.get("a", "")), str(lk.get("b", "")),
                     str(lk.get("type", "")), str(lk.get("net", "")) or "-"])
    if rows:
        tables.append(("互联链路表", ["A 端", "B 端", "类型", "网段"], rows, [5.0, 5.0, 2.5, 3.5]))

    # ACL 策略表
    rows = []
    for acl in acls:
        for r in acl.get("rules", []):
            rows.append([str(acl.get("id", "")), str(acl.get("type", "")),
                         str(acl.get("purpose", "")), str(r)])
    if rows:
        tables.append(("ACL 策略表", ["ACL ID", "类型", "用途", "规则"], rows, [2.0, 2.5, 4.5, 7.0]))

    return tables


def build_credentials(model, mask):
    """收集模型中的敏感凭据，供 --credentials 另存明文。"""
    creds = {}
    if model.get("meta", {}).get("credentials"):
        creds["meta.credentials"] = model["meta"]["credentials"]
    for dev in model.get("devices", []):
        hn = dev.get("hostname", "?")
        for k, v in dev.items():
            if isinstance(v, str) and any(k2 in k.lower() for k2 in SECRET_KEYS):
                creds.setdefault("设备 %s" % hn, {})[k] = v
    return creds


def write_credentials_docx(creds, out_path):
    Document, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT = ensure_docx()
    from docx.oxml.ns import qn as _qn
    qn = _qn
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    t = doc.add_heading(level=0)
    r = t.add_run("项目凭据清单（敏感 · 请加密保管）")
    set_cjk(r, "黑体")
    doc.add_paragraph("本文件含明文账号/密码/社区字符串，请勿明文外发，建议加密存储。")
    for group, items in creds.items():
        doc.add_heading(group, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["字段", "值"]):
            hdr[i].text = ""
            pr = hdr[i].paragraphs[0].add_run(col)
            set_cjk(pr, "黑体")
            pr.bold = True
        for k, v in items.items():
            cells = table.add_row().cells
            cells[0].text = ""
            cells[1].text = ""
            pr0 = cells[0].paragraphs[0].add_run(str(k))
            pr1 = cells[1].paragraphs[0].add_run(str(v))
            set_cjk(pr0, "宋体")
            set_cjk(pr1, "宋体")
        doc.add_paragraph("")
    doc.save(out_path)
    print("OK: 已生成明文凭据清单 -> %s（请加密保管）" % out_path)


def build(json_path, out_path, mask_secrets=False, no_toc=False, creds_path=None):
    Document, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT = ensure_docx()
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with open(json_path, "r", encoding="utf-8") as f:
        model = json.load(f)

    errs = validate(model)
    if errs:
        sys.stderr.write("ERROR: 项目模型校验失败：\n  - " + "\n  - ".join(errs) + "\n")
        sys.exit(3)

    # 明文凭据另存
    if creds_path:
        creds = build_credentials(model, mask_secrets)
        if creds:
            write_credentials_docx(creds, creds_path)

    doc = Document()

    # 默认正文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), "宋体")

    meta = model.get("meta", {})
    project_name = meta.get("project_name", "网络项目")
    author = meta.get("author", "AI")
    date = meta.get("date", "")

    # 标题
    t = doc.add_heading(level=0)
    r = t.add_run(project_name + " —— 网络项目部署设计文档")
    set_cjk(r, "黑体")

    # 目录（域）
    if not no_toc:
        doc.add_paragraph("目录")
        add_toc(doc)
        doc.add_page_break()

    # 基本信息
    info = doc.add_paragraph()
    for label, key in [("项目名称", "project_name"), ("编制", "author"),
                       ("日期", "date"), ("背景", "background"),
                       ("目标", "objective"), ("架构", "architecture")]:
        if key in meta and meta[key]:
            p = doc.add_paragraph()
            rr = p.add_run(f"{label}：")
            set_cjk(rr, "黑体")
            rb = p.add_run(meta[key])
            set_cjk(rb, "宋体")

    # 拓扑图
    doc.add_heading("一、网络拓扑图", level=1)
    topo = model.get("topology", "")
    if topo:
        para = doc.add_paragraph()
        run = para.add_run(topo)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rpr2 = run._element.get_or_add_rPr()
        rf2 = rpr2.find(qn('w:rFonts'))
        if rf2 is None:
            from docx.oxml import OxmlElement
            rf2 = OxmlElement('w:rFonts')
            rpr2.append(rf2)
        rf2.set(qn('w:eastAsia'), "宋体")
        rf2.set(qn('w:ascii'), "Consolas")
        rf2.set(qn('w:hAnsi'), "Consolas")

    # 台账表
    doc.add_heading("二、设备台账与配置总表", level=1)
    for idx, ledger in enumerate(model.get("ledgers", []), start=1):
        doc.add_heading(f"2.{idx} {ledger.get('title','台账')}", level=2)
        columns = ledger.get("columns", [])
        rows = ledger.get("rows", [])
        if not columns:
            continue
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = ""
            pr = hdr[i].paragraphs[0].add_run(col)
            set_cjk(pr, "黑体")
            pr.bold = True
            set_cell_shading(hdr[i], "D9E2F3")
        repeat_header_row(table)
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                v = mask_secret_value(val, columns[i]) if mask_secrets else val
                pr = cells[i].paragraphs[0].add_run(str(v))
                set_cjk(pr, "宋体")
        # 列宽自适应（按列数均分页宽约 16cm）
        set_table_col_widths(table, [16.0 / len(columns)] * len(columns))
        doc.add_paragraph("")

    # 关键决策记录（阶段3/4/5 确认项，必含）
    kd = model.get("key_decisions")
    if kd:
        doc.add_heading("三、关键决策记录", level=1)
        kd_cols = ["决策项", "确认选择", "理由/备注"]
        kd_table = doc.add_table(rows=1, cols=len(kd_cols))
        kd_table.style = "Table Grid"
        kd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, col in enumerate(kd_cols):
            kd_table.rows[0].cells[i].text = ""
            pr = kd_table.rows[0].cells[i].paragraphs[0].add_run(col)
            set_cjk(pr, "黑体")
            pr.bold = True
            set_cell_shading(kd_table.rows[0].cells[i], "D9E2F3")
        repeat_header_row(kd_table)
        for rec in kd:
            cells = kd_table.add_row().cells
            for i, key in enumerate(["decision", "choice", "rationale"]):
                cells[i].text = ""
                pr = cells[i].paragraphs[0].add_run(str(rec.get(key, "")))
                set_cjk(pr, "宋体")
        set_table_col_widths(kd_table, [4.0, 5.0, 7.0])
        doc.add_paragraph("")

    # 工程化规划表（自动生成）
    auto_tables = build_auto_tables(model, mask_secrets)
    if auto_tables:
        doc.add_heading("四、工程化规划表", level=1)
        for idx, (title, cols, rows, widths) in enumerate(auto_tables, start=1):
            doc.add_heading("4.%d %s" % (idx, title), level=2)
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, col in enumerate(cols):
                table.rows[0].cells[i].text = ""
                pr = table.rows[0].cells[i].paragraphs[0].add_run(col)
                set_cjk(pr, "黑体")
                pr.bold = True
                set_cell_shading(table.rows[0].cells[i], "D9E2F3")
            repeat_header_row(table)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    v = mask_secret_value(val, cols[i]) if mask_secrets else val
                    pr = cells[i].paragraphs[0].add_run(str(v))
                    set_cjk(pr, "宋体")
            set_table_col_widths(table, widths)
            doc.add_paragraph("")

    # 文字详述
    doc.add_heading("五、项目详细说明", level=1)
    for item in model.get("narrative", []):
        h = doc.add_heading(item.get("title", "说明"), level=2)
        for run in h.runs:
            set_cjk(run, "黑体")
        body = item.get("body", "")
        for line in body.split("\n"):
            p = doc.add_paragraph(line if line else " ")
            for run in p.runs:
                set_cjk(run, "宋体")

    # 落款
    doc.add_paragraph("")
    foot = doc.add_paragraph(f"—— 本文档由 network-deploy 技能生成，含账号密码等敏感信息，请妥善保管。")
    for run in foot.runs:
        set_cjk(run, "宋体")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 页码页脚
    add_page_number_footer(doc)

    out_dir = os.path.dirname(out_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)
    doc.save(out_path)
    print("OK: 已生成文档 -> %s" % out_path)
    if mask_secrets:
        print("INFO: --mask-secrets 已开启，文档中敏感值已打码为 ***。")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="网络项目部署文档生成器")
    ap.add_argument("json_path", nargs="?", default="项目模型.json", help="项目模型 JSON 路径")
    ap.add_argument("out_path", nargs="?", default="项目设计文档.docx", help="输出 docx 路径")
    ap.add_argument("--mask-secrets", action="store_true", help="文档中密码/密钥打码为 ***")
    ap.add_argument("--no-toc", action="store_true", help="不生成目录域")
    ap.add_argument("--credentials", default=None, help="明文凭据另存为独立 docx 路径")
    args = ap.parse_args()

    if not os.path.exists(args.json_path):
        sys.stderr.write("ERROR: 找不到 JSON 模型文件: %s\n" % args.json_path)
        sys.exit(1)
    build(args.json_path, args.out_path, mask_secrets=args.mask_secrets,
          no_toc=args.no_toc, creds_path=args.credentials)
